"""Parse CV from .docx file and extract text content."""

from pathlib import Path
from docx import Document

PROJECT_ROOT = Path(__file__).parent.parent
PROFILE_DIR = PROJECT_ROOT / "profile"


def parse_docx(filepath):
    """Extract all text from a .docx file."""
    doc = Document(filepath)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text.strip())

    return "\n".join(full_text)


def find_cv():
    """Find CV file in profile directory."""
    if not PROFILE_DIR.exists():
        return None

    for ext in ["*.docx", "*.doc"]:
        files = list(PROFILE_DIR.glob(ext))
        if files:
            return files[0]
    return None


if __name__ == "__main__":
    cv_path = find_cv()
    if cv_path:
        print(f"Found CV: {cv_path.name}\n")
        print("=" * 50)
        content = parse_docx(cv_path)
        print(content)
    else:
        print("No CV found in profile/ directory")
